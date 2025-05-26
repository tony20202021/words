#!/usr/bin/env python3
import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def generate_word_document(input_file, output_file=None, word_field_name="character", character_font_size=32, 
                       description_lines_count=3, description_limit=100,
                       transcription_font_size=12, default_font_size=11, page_margins=(1, 1, 1, 1), 
                       column_widths=None, character_line_spacing=24, transcription_limit=30):
    """
    Генерирует документ Microsoft Word из файла с данными о словах.
    
    Args:
        input_file (str): Путь к JSON-файлу с данными о словах
        output_file (str, optional): Путь для сохранения документа Word. 
                                     По умолчанию - то же имя файла с расширением .docx
        word_field_name (str, optional): Имя поля с основным словом. По умолчанию "character".
        character_font_size (int, optional): Размер шрифта для слов. По умолчанию 32.
        description_limit (int, optional): Лимит символов для колонки с описанием. По умолчанию 100.
        default_font_size (int, optional): Размер шрифта для остальных колонок. По умолчанию 11.
        page_margins (tuple, optional): Поля страницы (левое, правое, верхнее, нижнее) в дюймах.
        column_widths (dict, optional): Ширина колонок в процентах от ширины страницы.
        character_line_spacing (int, optional): Точное значение междустрочного интервала в пунктах для ячеек со словами.
        transcription_limit (int, optional): Лимит символов для колонки с транскрипцией. По умолчанию 30.
    
    Returns:
        str: Путь к созданному документу Word
    """
    # Проверка существования входного файла
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"Файл не найден: {input_file}")
    
    # Если выходной файл не указан, создаем имя на основе входного файла
    if output_file is None:
        output_file = os.path.splitext(input_file)[0] + ".docx"
    
    print(f"Чтение данных из файла: {input_file}")
    
    # Чтение JSON-файла
    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Создание списка записей и сортировка по частотности
    entries = []
    for key, value in data.items():
        # Проверяем, что частотность существует и является числом
        if 'frequency' in value and value['frequency']:
            try:
                frequency = int(value['frequency'])
                entries.append({
                    'key': key,
                    'frequency': frequency,
                    'character': value.get(word_field_name, ''),  # Используем word_field_name
                    'transcription': value.get('transcription', ''),
                    'description': value.get('description', [])
                })
            except (ValueError, TypeError):
                print(f"Пропуск записи {key}: некорректная частотность")
    
    # Сортировка по частотности (по возрастанию)
    entries.sort(key=lambda x: x['frequency'])
    
    print(f"Найдено {len(entries)} записей с корректной частотностью")
    
    # Создание нового документа Word
    doc = Document()
    
    # Установка полей страницы
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(page_margins[0])
        section.right_margin = Inches(page_margins[1])
        section.top_margin = Inches(page_margins[2])
        section.bottom_margin = Inches(page_margins[3])
    
    # Создание таблицы (изменено с 6 на 5 колонок, так как объединяем frequency и key)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Свойства для таблицы
    table.allow_autofit = False
    
    # Удаление первой строки (чтобы не было заголовка)
    table._element.remove(table.rows[0]._element)
    
    # Добавление данных в таблицу
    for entry in entries:
        # Добавление новой строки
        row = table.add_row().cells
        
        # Запрет разрыва строки таблицы между страницами
        row_element = row[0]._element.getparent()
        row_element.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit'] = '1'
        
        # 1. Частотность и номер (объединенные в одну колонку)
        freq_key_para = row[0].paragraphs[0]
        freq_key_text = f"{entry['frequency']}\n({entry['key']})"
        freq_key_run = freq_key_para.add_run(freq_key_text)
        freq_key_run.font.size = Pt(default_font_size)
        
        # 2. Слово/иероглиф (с увеличенным размером шрифта)
        char_para = row[1].paragraphs[0]
        char_run = char_para.add_run(entry['character'])
        char_run.font.size = Pt(character_font_size)
        
        # Установка точного межстрочного интервала
        char_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        char_para.paragraph_format.line_spacing = Pt(character_line_spacing)
        
        # 3. Транскрипция с ограничением длины
        transcription_text = entry['transcription']
        if len(transcription_text) > transcription_limit:
            transcription_text = transcription_text[:transcription_limit] + "..."
            
        trans_para = row[2].paragraphs[0]
        trans_run = trans_para.add_run(transcription_text)
        trans_run.font.size = Pt(transcription_font_size)
        
        # 4. Пустой столбец для комментариев
        comm_para = row[3].paragraphs[0]
        comm_run = comm_para.add_run("")
        comm_run.font.size = Pt(default_font_size)
        
        # 5. Перевод (описание) с ограничением длины
        description_text = ""
        if entry['description']:
            # Объединяем строки описания
            for desc in entry['description'][:description_lines_count]:
                if description_text:
                    description_text += "\n"
                
                # Обрезаем по лимиту символов
                if len(desc) > description_limit:
                    description_text += desc[:description_limit] + "..."
                else:
                    description_text += desc
        
        # Добавляем описание в документ
        desc_para = row[4].paragraphs[0]
        desc_run = desc_para.add_run(description_text)
        desc_run.font.size = Pt(default_font_size)
    
    # Задаем ширину колонок, если они указаны
    if column_widths:
        # Общая ширина таблицы (ориентировочно 6 дюймов для A4 с полями)
        total_width = 6  # в дюймах
        
        # Обновляем ключи колонок под новую структуру (без ключа 'key')
        column_keys = ['frequency', 'character', 'transcription', 'comments', 'description']
        for i, key in enumerate(column_keys):
            if key in column_widths:
                # Преобразуем процент в дюймы
                # Для первой колонки (frequency) теперь используем сумму процентов frequency и key
                width_percent = column_widths[key]
                if i == 0 and 'key' in column_widths:
                    width_percent = column_widths['frequency']  # Используем только значение для frequency
                
                width_inches = (width_percent / 100) * total_width
                for cell in table.columns[i].cells:
                    cell.width = Inches(width_inches)
    
    # Сохранение документа
    doc.save(output_file)
    
    print(f"Документ успешно создан: {output_file}")
    return output_file

def generate_excel_document(input_file, output_file=None, word_field_name="character", description_lines_count=3, 
                        description_limit=100, transcription_limit=30, column_widths=None):
    """
    Генерирует документ Microsoft Excel из файла с данными о словах.
    
    Args:
        input_file (str): Путь к JSON-файлу с данными о словах
        output_file (str, optional): Путь для сохранения документа Excel. 
                                     По умолчанию - то же имя файла с расширением .xlsx
        word_field_name (str, optional): Имя поля с основным словом. По умолчанию "character".
        description_lines_count (int, optional): Максимальное количество строк описания. По умолчанию 3.
        description_limit (int, optional): Лимит символов для колонки с описанием. По умолчанию 100.
        transcription_limit (int, optional): Лимит символов для колонки с транскрипцией. По умолчанию 30.
        column_widths (dict, optional): Ширина колонок в пикселях.
    
    Returns:
        str: Путь к созданному документу Excel
    """
    import pandas as pd
    import openpyxl
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.styles import Font, Alignment, PatternFill
    
    # Проверка существования входного файла
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"Файл не найден: {input_file}")
    
    # Если выходной файл не указан, создаем имя на основе входного файла
    if output_file is None:
        output_file = os.path.splitext(input_file)[0] + ".xlsx"
    
    print(f"Чтение данных из файла: {input_file}")
    
    # Чтение JSON-файла
    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Создание списка записей и сортировка по частотности
    entries = []
    for key, value in data.items():
        # Проверяем, что частотность существует и является числом
        if 'frequency' in value and value['frequency']:
            try:
                frequency = int(value['frequency'])
                
                # Подготовка описания
                description_text = ""
                if 'description' in value and value['description']:
                    # Объединяем строки описания
                    for desc in value['description'][:description_lines_count]:
                        if description_text:
                            description_text += "\n"
                        
                        # Обрезаем по лимиту символов
                        if len(desc) > description_limit:
                            description_text += desc[:description_limit] + "..."
                        else:
                            description_text += desc
                
                # Подготовка транскрипции
                transcription_text = value.get('transcription', '')
                if len(transcription_text) > transcription_limit:
                    transcription_text = transcription_text[:transcription_limit] + "..."
                
                entries.append({
                    'key': key,
                    'frequency': frequency,
                    word_field_name: value.get(word_field_name, ''),  # Используем word_field_name
                    'transcription': transcription_text,
                    'comments': '',  # Пустой столбец для комментариев
                    'description': description_text
                })
            except (ValueError, TypeError):
                print(f"Пропуск записи {key}: некорректная частотность")
    
    # Сортировка по частотности (по возрастанию)
    entries.sort(key=lambda x: x['frequency'])
    
    print(f"Найдено {len(entries)} записей с корректной частотностью")
    
    # Создание DataFrame
    df = pd.DataFrame(entries)
    
    # Создание Excel-файла
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Words"
    
    # Добавление заголовков
    headers = ['key', 'frequency', word_field_name, 'transcription', 'comments', 'description']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    
    # Добавление данных в Excel
    for i, entry in enumerate(entries, 2):  # начинаем со 2-й строки после заголовков
        ws.cell(row=i, column=1).value = entry['key']
        ws.cell(row=i, column=2).value = entry['frequency']
        
        # Настройка ячейки со словом
        char_cell = ws.cell(row=i, column=3)
        char_cell.value = entry[word_field_name]
        char_cell.font = Font(size=16)  # Увеличенный размер для слов
        char_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        ws.cell(row=i, column=4).value = entry['transcription']
        ws.cell(row=i, column=5).value = entry['comments']
        
        # Настройка ячейки с описанием
        desc_cell = ws.cell(row=i, column=6)
        desc_cell.value = entry['description']
        desc_cell.alignment = Alignment(wrap_text=True, vertical='top')
    
    # Настройка ширины колонок
    if column_widths:
        for i, header in enumerate(headers, 1):
            if header in column_widths:
                ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = column_widths[header]
    else:
        # Если ширина колонок не задана, устанавливаем приблизительные значения
        ws.column_dimensions['A'].width = 6     # key
        ws.column_dimensions['B'].width = 8     # frequency
        ws.column_dimensions['C'].width = 12    # word_field_name
        ws.column_dimensions['D'].width = 15    # transcription
        ws.column_dimensions['E'].width = 15    # comments
        ws.column_dimensions['F'].width = 45    # description
    
    # Автоматическая настройка высоты строк
    for i in range(2, len(entries) + 2):
        # Примерный расчет высоты строки в зависимости от количества строк в описании
        desc = entries[i-2]['description']
        num_lines = desc.count('\n') + 1
        row_height = max(14.5, num_lines * 14.5)  # Минимальная высота 14.5, для каждой строки добавляем 14.5
        ws.row_dimensions[i].height = row_height
    
    # Добавление автофильтра
    ws.auto_filter.ref = f"A1:F{len(entries) + 1}"
    
    # Закрепление заголовков
    ws.freeze_panes = 'A2'
    
    # Сохранение файла
    wb.save(output_file)
    
    print(f"Excel-документ успешно создан: {output_file}")
    return output_file


if __name__ == "__main__":
    # Задаем параметры как константы
    # INPUT_FILE = "chinese_characters_0_10000.json"
    INPUT_FILE = "chinese_characters_0_10000_description.valid.json.cleaned.json"
    OUTPUT_FILE_WORD = "chinese_characters_10_000.docx"
    OUTPUT_FILE_EXCEL = "chinese_characters_10_000.xlsx"

    # INPUT_FILE = "../data/fr.json.cleaned.json"
    # OUTPUT_FILE_WORD = "fr_10_000.docx"
    # OUTPUT_FILE_EXCEL = "fr_10_000.xlsx"

    # INPUT_FILE = "../data/eng.json.cleaned.json"
    # OUTPUT_FILE_WORD = "eng_10_000.docx"
    # OUTPUT_FILE_EXCEL = "eng_10_000.xlsx"

    # INPUT_FILE = "../data/deutch.json.cleaned.json"
    # OUTPUT_FILE_WORD = "deutch_10_000.docx"
    # OUTPUT_FILE_EXCEL = "deutch_10_000.xlsx"

    # INPUT_FILE = "../data/spain.json.cleaned.json"
    # OUTPUT_FILE_WORD = "spain_10_000.docx"
    # OUTPUT_FILE_EXCEL = "spain_10_000.xlsx"

    # Поле с основным словом
    WORD_FIELD_NAME = "word"

    # Размеры шрифтов
    CHARACTER_FONT_SIZE = 20
    TRANSCRIPTION_FONT_SIZE = 12
    DEFAULT_FONT_SIZE = 8
    
    # Лимит символов для описания
    DESCRIPTION_LINES_COUNT = 5
    DESCRIPTION_LIMIT = 285
    # Лимит символов для транскрипции
    TRANSCRIPTION_LIMIT = 115
        
    # Междустрочный интервал для ячеек с иероглифами (в пунктах)
    CHARACTER_LINE_SPACING = 22
    
    # Поля страницы в дюймах (левое, правое, верхнее, нижнее)
    PAGE_MARGINS = (0.3, 0.3, 0.1, 0.1)
    
    # Ширина колонок в процентах от ширины страницы для Word
    COLUMN_WIDTHS_WORD = {
        'frequency': 5,       # Частотность
        'character': 15,       # Иероглиф/Слово
        'transcription': 12,   # Транскрипция
        'comments': 12,        # Комментарии
        'description': 110      # Описание/перевод
    }
    
    # Ширина колонок в пикселях для Excel
    COLUMN_WIDTHS_EXCEL = {
        'key': 6,          # Ключ
        'frequency': 8,    # Частотность
        'character': 12,   # Иероглиф/Слово
        'transcription': 15, # Транскрипция
        'comments': 15,    # Комментарии
        'description': 145  # Описание/перевод
    }
        
    # Вызов функции генерации Word документа
    generate_word_document(
        INPUT_FILE, 
        OUTPUT_FILE_WORD,
        WORD_FIELD_NAME,  # Добавляем новый параметр
        CHARACTER_FONT_SIZE, 
        DESCRIPTION_LINES_COUNT,
        DESCRIPTION_LIMIT,
        TRANSCRIPTION_FONT_SIZE,
        DEFAULT_FONT_SIZE,
        PAGE_MARGINS,
        COLUMN_WIDTHS_WORD,
        CHARACTER_LINE_SPACING,
        TRANSCRIPTION_LIMIT
    )
    
    # Вызов функции генерации Excel документа
    generate_excel_document(
        INPUT_FILE,
        OUTPUT_FILE_EXCEL,
        WORD_FIELD_NAME,  # Добавляем новый параметр
        DESCRIPTION_LINES_COUNT,
        DESCRIPTION_LIMIT,
        TRANSCRIPTION_LIMIT,
        COLUMN_WIDTHS_EXCEL
    )
