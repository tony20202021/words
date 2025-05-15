import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_word_document_with_modified_table(json_path, output_docx_path, max_cells=214, full_image_scale=0.4, 
                                  character_image_scale=0.6, region_image_scale=0.8, inner_table_height=3000,
                                  image_column_width=8):
    """
    Создает документ Word с таблицей, содержащей данные из JSON-файла и изображения.
    Каждая пара ячеек (четная и нечетная) размещаются в двух строках таблицы.
    Каждая пара имеет общую ячейку полного изображения.
    
    Параметры:
    json_path (str): Путь к JSON-файлу с информацией о ячейках
    output_docx_path (str): Путь для сохранения документа Word
    max_cells (int): Максимальное количество ячеек (строк таблицы)
    full_image_scale (float): Коэффициент масштабирования полных изображений ячеек
    character_image_scale (float): Коэффициент масштабирования изображений иероглифов
    region_image_scale (float): Коэффициент масштабирования других изображений регионов
    inner_table_height (int): Высота строки внутренней таблицы в твипах (1440 = 1 дюйм)
    image_column_width (float): Ширина колонки изображений в сантиметрах
    """
    # Загружаем данные из JSON-файла
    print(f"Загрузка данных из {json_path}...")
    with open(json_path, 'r', encoding='utf-8') as f:
        cells_info = json.load(f)
    
    # Получаем список ключей, отсортированных по номеру ячейки
    sorted_keys = sorted([int(k) for k in cells_info.keys()])
    sorted_keys = [str(k) for k in sorted_keys if int(k) <= max_cells]
    
    print(f"Создание таблицы для {len(sorted_keys)} ячеек...")
    print(f"Масштаб изображений:")
    print(f" - Полные ячейки: {full_image_scale * 100}%")
    print(f" - Иероглифы: {character_image_scale * 100}%")
    print(f" - Пиньинь и перевод: {region_image_scale * 100}%")
    print(f"Высота строки внутренней таблицы: {inner_table_height} твипов (~{inner_table_height/1440:.2f} дюймов)")
    print(f"Ширина колонки изображений: {image_column_width} см")
    
    # Создаем новый документ Word
    doc = Document()
    
    # Настраиваем поля документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # Функция для очистки содержимого ячейки и добавления нового текста без лишних переводов строк
    def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        # Удаляем все параграфы в ячейке
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        
        # Добавляем новый параграф с нужным текстом
        paragraph = cell.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        return paragraph
    
    # Функция для установки минимальной высоты строки
    def set_table_height(table, row_idx, height):
        tr = table.rows[row_idx]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
    
    # Функция для предотвращения разрыва строки между страницами
    def prevent_row_break(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
    
    # Рассчитываем ширину изображений на основе коэффициентов масштабирования
    image_widths = {
        'full': Cm(3.8 * full_image_scale),
        'character': Cm(3.8 * character_image_scale),
        'pinyin': Cm(2.8 * region_image_scale),
        'translation': Cm(3.8 * region_image_scale)
    }
    
    # Создаем таблицу без заголовка
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    
    # Устанавливаем ширину столбцов
    column_widths = [1, image_column_width, 4, 3, 2.5, 4]
    
    # Заполняем таблицу данными
    for pair_idx in range(len(sorted_keys) // 2):
        # Индексы ячеек в паре
        idx1 = pair_idx * 2
        idx2 = idx1 + 1
        
        # Ключи ячеек в паре
        key1 = sorted_keys[idx1]
        key2 = sorted_keys[idx2] if idx2 < len(sorted_keys) else None
        
        # Если не хватает ячеек для полной пары, пропускаем
        if key2 is None:
            continue
        
        # Получаем информацию о ячейках
        cell_info1 = cells_info[key1]
        cell_info2 = cells_info[key2]
        
        # Добавляем две строки для этой пары
        row1 = table.add_row()
        row2 = table.add_row()
        
        # Устанавливаем свойство не разрывать строки между страницами
        prevent_row_break(row1)
        prevent_row_break(row2)
        
        # Настраиваем ширину колонок для этих строк
        for j, width in enumerate(column_widths):
            row1.cells[j].width = Cm(width)
            row2.cells[j].width = Cm(width)
        
        # Заполняем номера ячеек без лишних переводов строк
        set_cell_text(row1.cells[0], key1)
        set_cell_text(row2.cells[0], key2)
        
        # Получаем ячейку для размещения внутренней таблицы
        img_cell = row1.cells[1]
        img_cell.merge(row2.cells[1])  # Объединяем ячейки по вертикали
        
        # Очищаем содержимое объединенной ячейки
        for p in img_cell.paragraphs:
            p._element.getparent().remove(p._element)
        
        # Создаем внутреннюю таблицу с 1 строкой и 2 колонками
        inner_table = img_cell.add_table(rows=1, cols=2)
        inner_table.style = 'Table Grid'
        inner_table.autofit = False
        
        # Устанавливаем высоту строки во внутренней таблице из параметра
        set_table_height(inner_table, 0, inner_table_height)
        
        # Заполняем внутреннюю таблицу изображениями
        inner_cell1 = inner_table.cell(0, 0)
        inner_p1 = inner_cell1.paragraphs[0]
        inner_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем первое изображение
        full_image_path1 = cell_info1['image_path']
        if os.path.exists(full_image_path1):
            run1 = inner_p1.add_run()
            run1.add_picture(full_image_path1, width=image_widths['full'])
        else:
            inner_p1.add_run(f"Ячейка {key1} (файл не найден)")
        
        # Устанавливаем выравнивание по верху для первой ячейки
        inner_cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        inner_cell2 = inner_table.cell(0, 1)
        inner_p2 = inner_cell2.paragraphs[0]
        inner_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем второе изображение
        full_image_path2 = cell_info2['image_path']
        if os.path.exists(full_image_path2):
            run2 = inner_p2.add_run()
            run2.add_picture(full_image_path2, width=image_widths['full'])
        else:
            inner_p2.add_run(f"Ячейка {key2} (файл не найден)")
        
        # Устанавливаем выравнивание по низу для второй ячейки
        inner_cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        
        # Заполняем остальные ячейки для первой строки в паре
        # Иероглиф
        character_image_path1 = cell_info1['regions']['character']['image_path']
        if os.path.exists(character_image_path1):
            p_ierogl1 = set_cell_text(row1.cells[2], "")
            run_ierogl1 = p_ierogl1.add_run()
            run_ierogl1.add_picture(character_image_path1, width=image_widths['character'])
        else:
            set_cell_text(row1.cells[2], f"Файл не найден: {character_image_path1}")
        
        # Транскрипция (пиньинь)
        pinyin_image_path1 = cell_info1['regions']['pinyin']['image_path']
        if os.path.exists(pinyin_image_path1):
            p_pinyin1 = set_cell_text(row1.cells[3], "")
            run_pinyin1 = p_pinyin1.add_run()
            run_pinyin1.add_picture(pinyin_image_path1, width=image_widths['pinyin'])
        else:
            set_cell_text(row1.cells[3], f"Файл не найден: {pinyin_image_path1}")
        
        # Комментарии (оставляем пустыми)
        set_cell_text(row1.cells[4], "")
        
        # Перевод
        translation_image_path1 = cell_info1['regions']['translation']['image_path']
        if os.path.exists(translation_image_path1):
            p_transl1 = set_cell_text(row1.cells[5], "")
            run_transl1 = p_transl1.add_run()
            run_transl1.add_picture(translation_image_path1, width=image_widths['translation'])
        else:
            set_cell_text(row1.cells[5], f"Файл не найден: {translation_image_path1}")
        
        # Заполняем остальные ячейки для второй строки в паре
        # Иероглиф
        character_image_path2 = cell_info2['regions']['character']['image_path']
        if os.path.exists(character_image_path2):
            p_ierogl2 = set_cell_text(row2.cells[2], "")
            run_ierogl2 = p_ierogl2.add_run()
            run_ierogl2.add_picture(character_image_path2, width=image_widths['character'])
        else:
            set_cell_text(row2.cells[2], f"Файл не найден: {character_image_path2}")
        
        # Транскрипция (пиньинь)
        pinyin_image_path2 = cell_info2['regions']['pinyin']['image_path']
        if os.path.exists(pinyin_image_path2):
            p_pinyin2 = set_cell_text(row2.cells[3], "")
            run_pinyin2 = p_pinyin2.add_run()
            run_pinyin2.add_picture(pinyin_image_path2, width=image_widths['pinyin'])
        else:
            set_cell_text(row2.cells[3], f"Файл не найден: {pinyin_image_path2}")
        
        # Комментарии (оставляем пустыми)
        set_cell_text(row2.cells[4], "")
        
        # Перевод
        translation_image_path2 = cell_info2['regions']['translation']['image_path']
        if os.path.exists(translation_image_path2):
            p_transl2 = set_cell_text(row2.cells[5], "")
            run_transl2 = p_transl2.add_run()
            run_transl2.add_picture(translation_image_path2, width=image_widths['translation'])
        else:
            set_cell_text(row2.cells[5], f"Файл не найден: {translation_image_path2}")
        
        # Выводим прогресс
        if (pair_idx + 1) % 5 == 0 or pair_idx == (len(sorted_keys) // 2) - 1:
            print(f"Обработано {(pair_idx + 1) * 2} из {len(sorted_keys)} ячеек...")
    
    # Обрабатываем последнюю ячейку, если общее количество нечетное
    if len(sorted_keys) % 2 != 0:
        last_idx = len(sorted_keys) - 1
        last_key = sorted_keys[last_idx]
        last_cell_info = cells_info[last_key]
        
        # Добавляем строку для последней ячейки
        last_row = table.add_row()
        
        # Устанавливаем свойство не разрывать строку между страницами
        prevent_row_break(last_row)
        
        # Настраиваем ширину колонок
        for j, width in enumerate(column_widths):
            last_row.cells[j].width = Cm(width)
        
        # Заполняем номер ячейки
        set_cell_text(last_row.cells[0], last_key)
        
        # Полное изображение
        full_image_path = last_cell_info['image_path']
        if os.path.exists(full_image_path):
            p_full = set_cell_text(last_row.cells[1], "")
            run_full = p_full.add_run()
            run_full.add_picture(full_image_path, width=image_widths['full'])
        else:
            set_cell_text(last_row.cells[1], f"Файл не найден: {full_image_path}")
        
        # Иероглиф
        character_image_path = last_cell_info['regions']['character']['image_path']
        if os.path.exists(character_image_path):
            p_ierogl = set_cell_text(last_row.cells[2], "")
            run_ierogl = p_ierogl.add_run()
            run_ierogl.add_picture(character_image_path, width=image_widths['character'])
        else:
            set_cell_text(last_row.cells[2], f"Файл не найден: {character_image_path}")
        
        # Транскрипция (пиньинь)
        pinyin_image_path = last_cell_info['regions']['pinyin']['image_path']
        if os.path.exists(pinyin_image_path):
            p_pinyin = set_cell_text(last_row.cells[3], "")
            run_pinyin = p_pinyin.add_run()
            run_pinyin.add_picture(pinyin_image_path, width=image_widths['pinyin'])
        else:
            set_cell_text(last_row.cells[3], f"Файл не найден: {pinyin_image_path}")
        
        # Комментарии (оставляем пустыми)
        set_cell_text(last_row.cells[4], "")
        
        # Перевод
        translation_image_path = last_cell_info['regions']['translation']['image_path']
        if os.path.exists(translation_image_path):
            p_transl = set_cell_text(last_row.cells[5], "")
            run_transl = p_transl.add_run()
            run_transl.add_picture(translation_image_path, width=image_widths['translation'])
        else:
            set_cell_text(last_row.cells[5], f"Файл не найден: {translation_image_path}")
    
    # Сохраняем документ
    print(f"Сохранение документа в {output_docx_path}...")
    doc.save(output_docx_path)
    print(f"Документ успешно создан!")


if __name__ == "__main__":
    print("Создание документа Word с модифицированной структурой таблицы...")

    output_dir = "output_cells"
    json_path = os.path.join(output_dir, "cells_info.json")

    create_word_document_with_modified_table(
        json_path, 
        'chinese_radicals_table.docx', 
        max_cells=214,
        full_image_scale=0.25,
        character_image_scale=0.3,
        region_image_scale=0.55,
        inner_table_height=780, 
        image_column_width=4
    )
