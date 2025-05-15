import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

def create_sample_images(count=10):
    """
    Создает несколько тестовых изображений для примера.
    
    Параметры:
        count (int): Количество пар изображений для создания
    
    Возвращает:
        list: Список пар путей к изображениям [(путь1_1, путь1_2), (путь2_1, путь2_2), ...]
    """
    # Создаем папку для изображений, если её нет
    images_dir = "sample_images"
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)
    
    # Список для хранения путей к изображениям
    image_pairs = []
    
    # Создаем пары изображений
    for i in range(1, count + 1):
        # Пути к файлам изображений для пары
        image1_path = os.path.join(images_dir, f"sample_image{i}_1.png")
        image2_path = os.path.join(images_dir, f"sample_image{i}_2.png")
        
        # Первое изображение - красный оттенок с номером ячейки
        red_value = min(255, 100 + (i * 15))  # Различные оттенки красного
        img1 = Image.new('RGB', (200, 200), color=(red_value, 100, 100))
        draw1 = ImageDraw.Draw(img1)
        
        # Попытка использовать системный шрифт или создать текст без шрифта
        try:
            font = ImageFont.truetype("arial.ttf", 20)
            draw1.text((40, 90), f"Ячейка {i*2-1}", font=font, fill=(255, 255, 255))
        except:
            draw1.text((40, 90), f"Ячейка {i*2-1}", fill=(255, 255, 255))
        
        # Сохраняем первое изображение
        img1.save(image1_path)
        
        # Второе изображение - синий оттенок с номером ячейки
        blue_value = min(255, 100 + (i * 15))  # Различные оттенки синего
        img2 = Image.new('RGB', (200, 200), color=(100, 100, blue_value))
        draw2 = ImageDraw.Draw(img2)
        
        # Попытка использовать системный шрифт или создать текст без шрифта
        try:
            font = ImageFont.truetype("arial.ttf", 20)
            draw2.text((40, 90), f"Ячейка {i*2}", font=font, fill=(255, 255, 255))
        except:
            draw2.text((40, 90), f"Ячейка {i*2}", fill=(255, 255, 255))
        
        # Сохраняем второе изображение
        img2.save(image2_path)
        
        # Добавляем пару путей в список
        image_pairs.append((image1_path, image2_path))
    
    return image_pairs

def test_multi_row_table(rows=10, output_filename="multi_row_table.docx", inner_table_height=2000, 
                        image_width=3):
    """
    Создает тестовую таблицу с указанным количеством строк и изображениями.
    Исправлена проблема с лишними переводами строк в заголовках и номерах ячеек.
    
    Параметры:
    rows (int): Количество строк в таблице (пар ячеек)
    output_filename (str): Имя выходного файла
    inner_table_height (int): Высота строки внутренней таблицы в твипах (1440 = 1 дюйм)
    image_width (float): Ширина изображений в сантиметрах
    """
    # Создаем тестовые изображения для каждой пары строк
    image_pairs = create_sample_images(rows)
    
    doc = Document()
    
    # Настраиваем поля документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # Создаем таблицу с 6 столбцами и необходимым количеством строк
    table = doc.add_table(rows=rows*2, cols=6)
    table.style = 'Table Grid'
    
    # Функция для очистки содержимого ячейки и добавления нового текста без лишних переводов строк
    def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        # Удаляем все параграфы в ячейке
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        
        # Добавляем новый параграф с нужным текстом
        paragraph = cell.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        return paragraph
    
    # Устанавливаем ширину столбцов
    column_widths = [1, 4, 4, 3, 2.5, 4]  # Относительная ширина столбцов
    
    # Настраиваем ширину колонок для всех ячеек
    for i in range(rows*2):  # для всех строк
        for j, width in enumerate(column_widths):
            table.cell(i, j).width = Cm(width)
    
    # Добавляем заголовки таблицы (без лишних переводов строк)
    headers = ["№", "Изображение", "Иероглиф", "Транскрипция", "Комментарии", "Перевод"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header)
    
    # Функция для установки минимальной высоты строки
    def set_table_height(table, row_idx, height):
        tr = table.rows[row_idx]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
    
    # Заполняем таблицу данными
    for row_idx in range(rows):
        # Получаем пару изображений для текущей пары строк
        image1_path, image2_path = image_pairs[row_idx]
        
        # Индексы строк в таблице
        idx1 = row_idx * 2
        idx2 = idx1 + 1
        
        # Заполняем номера ячеек без лишних переводов строк
        set_cell_text(table.cell(idx1, 0), str(idx1 + 1))
        set_cell_text(table.cell(idx2, 0), str(idx2 + 1))
        
        # Получаем ячейку для размещения внутренней таблицы
        img_cell = table.cell(idx1, 1)
        img_cell.merge(table.cell(idx2, 1))  # Объединяем ячейки по вертикали
        
        # Очищаем содержимое объединенной ячейки
        for p in img_cell.paragraphs:
            p._element.getparent().remove(p._element)
        
        # Создаем внутреннюю таблицу с 1 строкой и 2 колонками
        inner_table = img_cell.add_table(rows=1, cols=2)
        inner_table.style = 'Table Grid'
        inner_table.autofit = False
        
        # Устанавливаем высоту строки во внутренней таблице из параметра
        set_table_height(inner_table, 0, inner_table_height)
        
        # Заполняем внутреннюю таблицу изображениями
        inner_cell1 = inner_table.cell(0, 0)
        inner_p1 = inner_cell1.paragraphs[0]
        inner_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем первое изображение с шириной из параметра
        if os.path.exists(image1_path):
            run1 = inner_p1.add_run()
            run1.add_picture(image1_path, width=Cm(image_width))
        else:
            inner_p1.add_run(f"Ячейка {idx1+1} (файл не найден)")
        
        # Устанавливаем выравнивание по верху для первой ячейки
        inner_cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        inner_cell2 = inner_table.cell(0, 1)
        inner_p2 = inner_cell2.paragraphs[0]
        inner_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем второе изображение с шириной из параметра
        if os.path.exists(image2_path):
            run2 = inner_p2.add_run()
            run2.add_picture(image2_path, width=Cm(image_width))
        else:
            inner_p2.add_run(f"Ячейка {idx2+1} (файл не найден)")
        
        # Устанавливаем выравнивание по низу для второй ячейки
        inner_cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        
        # Заполняем остальные ячейки для первой строки в паре (без лишних переводов строк)
        # Иероглиф
        set_cell_text(table.cell(idx1, 2), f"一 {idx1+1}")
        
        # Транскрипция
        set_cell_text(table.cell(idx1, 3), f"yī-{idx1+1}")
        
        # Комментарии (оставляем пустыми)
        set_cell_text(table.cell(idx1, 4), "")
        
        # Перевод
        set_cell_text(table.cell(idx1, 5), f"единица-{idx1+1}")
        
        # Заполняем остальные ячейки для второй строки в паре (без лишних переводов строк)
        # Иероглиф
        set_cell_text(table.cell(idx2, 2), f"丨 {idx2+1}")
        
        # Транскрипция
        set_cell_text(table.cell(idx2, 3), f"gŭn-{idx2+1}")
        
        # Комментарии (оставляем пустыми)
        set_cell_text(table.cell(idx2, 4), "")
        
        # Перевод
        set_cell_text(table.cell(idx2, 5), f"вертикальная-{idx2+1}")
    
    # Сохраняем документ
    doc.save(output_filename)
    print(f"Таблица с {rows} парами строк создана в файле {output_filename}")
    print(f"Высота строки внутренней таблицы: {inner_table_height} твипов (~{inner_table_height/1440:.2f} дюймов)")
    print(f"Ширина изображений: {image_width} см")


if __name__ == "__main__":
    # Создаем таблицу с 10 строками (20 ячеек)
    test_multi_row_table(rows=10, inner_table_height=4000)